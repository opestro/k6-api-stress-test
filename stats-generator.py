import json
import os
from docx import Document

def extract_metrics(filename):
    if not os.path.isfile(filename):
        print(f"Error: File '{filename}' not found.")
        return None
    
    metrics = {
        'http_req_duration': {'max': 0},
        'http_reqs': {'count': 0},
        'http_req_failed': {'fails': 0}
    }
    
    try:
        with open(filename, 'r') as f:
            for line in f:
                try:
                    data = json.loads(line.strip())
                    
                    metric_type = data.get('metric')
                    if metric_type:
                        if metric_type == 'http_reqs':
                            # Add debug information
                            print(f"Processing http_reqs: {data}")
                            value = data['data'].get('value', 0)
                            metrics['http_reqs']['count'] += value
                        elif metric_type == 'http_req_duration':
                            # Add debug information
                            print(f"Processing http_req_duration: {data}")
                            value = data['data'].get('value', 0)
                            if value > metrics['http_req_duration']['max']:
                                metrics['http_req_duration']['max'] = value
                        elif metric_type == 'http_req_failed':
                            # Add debug information
                            print(f"Processing http_req_failed: {data}")
                            value = data['data'].get('value', 0)
                            metrics['http_req_failed']['fails'] += value
                    
                except json.JSONDecodeError as e:
                    print(f"Error decoding JSON line: {e}")
    
    except IOError as e:
        print(f"Error reading file: {e}")
        return None

    return (
        metrics['http_req_duration']['max'],
        metrics['http_reqs']['count'],
        metrics['http_req_failed']['fails']
    )

# Create a new Word document
doc = Document()
doc.add_heading('K6 Stress Test Report', 0)

# Add a summary section
doc.add_heading('Test Summary', level=1)

# Create a table for the test results
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Stage'
hdr_cells[1].text = 'Max Response Time (ms)'
hdr_cells[2].text = 'Success Requests'
hdr_cells[3].text = 'Error Requests'

# Process each stage and add rows to the table
stages = [
    ('20 VUs', 'stage_20_vus_output.json'),
    ('100 VUs', 'stage_100_vus_output.json'),
    ('500 VUs', 'stage_500_vus_output.json'),
    ('1000 VUs', 'stage_1000_vus_output.json')
]

for stage_name, file_name in stages:
    metrics = extract_metrics(file_name)
    if metrics:
        max_response_time, successful_requests, error_requests = metrics
        row_cells = table.add_row().cells
        row_cells[0].text = stage_name
        row_cells[1].text = str(max_response_time)
        row_cells[2].text = str(successful_requests)
        row_cells[3].text = str(error_requests)

# Add a section for more details if needed
doc.add_heading('Details', level=1)
doc.add_paragraph('This report includes the essential metrics from the K6 stress test for each stage of the test.')

# Save the document to a file
doc.save('da3em_Stress_Test_Report_Combined.docx')
